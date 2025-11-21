"""
Document processing utilities for various file formats.
Adaptado de doc-reader-main para balcon_demo_local.
"""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

import pypdf
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from django.conf import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document loading, processing, and chunking."""

    def __init__(self):
        # Configuración desde settings de Django
        chunk_size = getattr(settings, 'RAG_CHUNK_SIZE', 1024)
        chunk_overlap = getattr(settings, 'RAG_CHUNK_OVERLAP', 512)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        self.max_file_size_mb = getattr(settings, 'RAG_MAX_FILE_SIZE_MB', 50)
        self.supported_formats = ['pdf', 'docx', 'txt', 'md']

    def load_document(self, file_path: str) -> str:
        """Load and extract text from various document formats."""
        path_obj = Path(file_path)

        if not path_obj.exists():
            raise FileNotFoundError(f"Documento no encontrado: {path_obj}")

        # Validar tamaño del archivo
        file_size_mb = path_obj.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValueError(
                f"Documento muy grande: {file_size_mb:.2f}MB (máximo: {self.max_file_size_mb}MB)"
            )

        file_extension = path_obj.suffix.lower().lstrip('.')

        if file_extension not in self.supported_formats:
            raise ValueError(
                f"Formato no soportado: {file_extension}. "
                f"Formatos válidos: {', '.join(self.supported_formats)}"
            )

        logger.info(f"📄 Cargando documento: {path_obj.name}")

        if file_extension == 'pdf':
            return self._load_pdf(path_obj)
        elif file_extension == 'docx':
            return self._load_docx(path_obj)
        elif file_extension in ['txt', 'md']:
            return self._load_text(path_obj)
        else:
            raise ValueError(f"Handler no implementado para: {file_extension}")

    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Página {page_num + 1}/{total_pages} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(
                            f"⚠️ Error extrayendo página {page_num + 1} de {file_path.name}: {e}"
                        )
                        continue
                        
            if not text.strip():
                raise ValueError("El PDF no contiene texto extraíble")
                
        except Exception as e:
            logger.error(f"❌ Error cargando PDF {file_path.name}: {e}")
            raise

        return text.strip()

    def _load_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(str(file_path))
            text = ""

            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            if not text.strip():
                raise ValueError("El DOCX está vacío o no contiene texto")

            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ Error cargando DOCX {file_path.name}: {e}")
            raise

    def _load_text(self, file_path: Path) -> str:
        """Load plain text files (TXT, MD)."""
        try:
            # Intentar con UTF-8 primero
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read().strip()
        except UnicodeDecodeError:
            # Fallback a latin-1 para archivos antiguos
            logger.info(f"📝 Usando encoding latin-1 para {file_path.name}")
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read().strip()
        
        if not text:
            raise ValueError("El archivo está vacío")
            
        return text

    def chunk_document(
        self, 
        text: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Split document into chunks for vector storage."""
        if metadata is None:
            metadata = {}

        chunks = self.text_splitter.split_text(text)

        if not chunks:
            raise ValueError("No se pudieron generar chunks del documento")

        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **metadata,
                "chunk_id": i,
                "chunk_size": len(chunk),
                "total_chunks": len(chunks)
            }
            documents.append(Document(
                page_content=chunk,
                metadata=chunk_metadata
            ))

        logger.info(f"✅ Creados {len(documents)} chunks del documento")
        return documents

    def process_document(
        self, 
        file_path: str, 
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Complete document processing pipeline."""
        path_obj = Path(file_path)

        # 1. Cargar y extraer texto
        text = self.load_document(file_path)

        # 2. Preparar metadata enriquecida
        metadata = {
            "source": str(path_obj),
            "filename": path_obj.name,
            "file_type": path_obj.suffix.lower().lstrip('.'),
            "file_size": path_obj.stat().st_size,
            "word_count": len(text.split())
        }

        # 3. Agregar metadata adicional (categoria, role_filter, etc.)
        if additional_metadata:
            metadata.update(additional_metadata)

        # 4. Generar chunks con metadata
        documents = self.chunk_document(text, metadata)

        logger.info(
            f"📦 Procesado {path_obj.name}: "
            f"{metadata['word_count']} palabras → {len(documents)} chunks"
        )

        return documents
